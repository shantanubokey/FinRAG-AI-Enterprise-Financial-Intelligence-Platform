"""DOCX loader using python-docx. Extracts text, tables, and heading structure."""

from typing import Any

from config.logging_config import get_logger
from ingestion.loaders.base_loader import BaseDocumentLoader, RawDocument

logger = get_logger(__name__)


class DOCXLoader(BaseDocumentLoader):

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def load(self, content: bytes, filename: str) -> RawDocument:
        import io
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(io.BytesIO(content))
        paragraphs: list[str] = []
        tables: list[dict[str, Any]] = []
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Track section headings for metadata
            if para.style.name.startswith("Heading"):
                current_section = text
            paragraphs.append(text)

        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append({
                "index": i,
                "headers": rows[0] if rows else [],
                "rows": rows[1:] if len(rows) > 1 else [],
                "data": rows,
            })

        return RawDocument(
            content="\n\n".join(paragraphs),
            tables=tables,
            metadata={"filename": filename, "paragraph_count": len(paragraphs)},
            page_contents=paragraphs,
        )
